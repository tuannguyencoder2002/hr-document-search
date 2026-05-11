"""Generate the project report as a .docx file.

Run: python scripts/generate_report.py
Output: report/Document_Search_Assistant_Report.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "report"
OUTPUT_DIR.mkdir(exist_ok=True)
OUTPUT_FILE = OUTPUT_DIR / "Document_Search_Assistant_Report.docx"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table


def build_report():
    doc = Document()
    set_normal_style(doc)

    # === TRANG BÌA ===
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO ĐỒ ÁN")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Document Search Assistant\nHệ thống RAG tìm kiếm tài liệu chạy 100% local")
    run.font.size = Pt(16)

    for _ in range(3):
        doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Sinh viên thực hiện: Nguyễn Quý Tuấn\n").font.size = Pt(13)
    info.add_run("Năm: 2026").font.size = Pt(13)

    doc.add_page_break()

    # === MỤC LỤC ===
    add_heading_styled(doc, "Mục lục", level=1)
    toc_items = [
        "I. Dẫn nhập",
        "II. Cơ sở lý thuyết về RAG",
        "   II.1. Retrieval-Augmented Generation là gì?",
        "   II.2. Kiến trúc RAG hiện đại",
        "   II.3. Hybrid Search + Reranking",
        "III. Kiến trúc hệ thống",
        "   III.1. Tổng quan kiến trúc",
        "   III.2. Ingestion Pipeline",
        "   III.3. Search + Generation Pipeline",
        "   III.4. Image Search (CLIP)",
        "IV. Công nghệ sử dụng",
        "   IV.1. Bảng công nghệ",
        "   IV.2. Giải thích lựa chọn model",
        "V. Giao diện người dùng",
        "VI. Kết quả và đánh giá",
        "VII. Kết luận và hướng phát triển",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number" if not item.startswith("   ") else "List Bullet")
    doc.add_page_break()

    # === BẢNG THUẬT NGỮ ===
    add_heading_styled(doc, "Bảng thuật ngữ", level=1)
    terms = [
        ("RAG", "Retrieval-Augmented Generation — kỹ thuật bổ sung ngữ cảnh từ nguồn dữ liệu bên ngoài cho LLM"),
        ("LLM", "Large Language Model — mô hình ngôn ngữ lớn (GPT, Qwen, Llama...)"),
        ("Embedding", "Biểu diễn văn bản/ảnh dưới dạng vector số thực trong không gian n-chiều"),
        ("Hybrid Search", "Kết hợp tìm kiếm ngữ nghĩa (dense) và từ khóa (sparse)"),
        ("RRF", "Reciprocal Rank Fusion — thuật toán gộp kết quả từ nhiều nguồn xếp hạng"),
        ("Cross-Encoder", "Mô hình đánh giá mức độ liên quan bằng cách đọc đồng thời query + document"),
        ("Chunking", "Chia nhỏ văn bản dài thành các đoạn ngắn để tối ưu embedding"),
        ("CLIP", "Contrastive Language-Image Pre-training — mô hình liên kết text ↔ image"),
        ("Qdrant", "Vector database hỗ trợ hybrid search (dense + sparse)"),
        ("Ollama", "Runtime chạy LLM local trên GPU consumer"),
    ]
    add_table(doc, ["Thuật ngữ", "Mô tả"], terms)
    doc.add_page_break()

    # === I. DẪN NHẬP ===
    add_heading_styled(doc, "I. Dẫn nhập", level=1)
    doc.add_paragraph(
        "Sự bùng nổ của các Large Language Models (LLMs) như ChatGPT, Gemini, Claude "
        "đã định hình lại lĩnh vực NLP. Tuy nhiên, các mô hình này vẫn đối mặt với "
        "những hạn chế cố hữu: tri thức bị giới hạn tại thời điểm huấn luyện, hiện "
        "tượng hallucination, và thiếu hụt kiến thức về dữ liệu riêng tư."
    )
    doc.add_paragraph(
        "Để giải quyết vấn đề này, kỹ thuật Retrieval-Augmented Generation (RAG) cho "
        "phép LLMs tiếp cận nguồn dữ liệu bên ngoài mà không cần fine-tuning. Bài báo "
        "cáo này trình bày hệ thống Document Search Assistant — một RAG system chạy "
        "hoàn toàn local trên GPU consumer (RTX 5060, 8GB VRAM), phục vụ tìm kiếm và "
        "hỏi đáp trên tài liệu học tập (PDF, DOCX, TXT)."
    )
    doc.add_paragraph(
        "Điểm nổi bật của hệ thống:\n"
        "• Hybrid Search (dense + sparse) với RRF fusion\n"
        "• Cross-encoder reranking cho độ chính xác cao\n"
        "• Image search bằng CLIP (tìm ảnh trong tài liệu bằng ảnh)\n"
        "• Streaming LLM response (Qwen3-4B)\n"
        "• 100% offline, dữ liệu không rời máy\n"
        "• Giao diện Next.js 14 hiện đại với inline PDF viewer"
    )
    doc.add_page_break()

    # === II. CƠ SỞ LÝ THUYẾT ===
    add_heading_styled(doc, "II. Cơ sở lý thuyết về RAG", level=1)

    add_heading_styled(doc, "II.1. Retrieval-Augmented Generation là gì?", level=2)
    doc.add_paragraph(
        "RAG (Lewis et al., 2020) là kỹ thuật kết hợp hai thành phần:\n"
        "1. Retriever: tìm kiếm các đoạn văn bản liên quan từ kho dữ liệu\n"
        "2. Generator: LLM sinh câu trả lời dựa trên ngữ cảnh được truy xuất\n\n"
        "Thay vì fine-tune LLM trên dữ liệu mới (tốn kém, mất tính tổng quát), "
        "RAG giữ nguyên trọng số LLM và chỉ cung cấp thông tin liên quan qua prompt. "
        "Đây là hướng tiếp cận In-Context Learning."
    )

    add_heading_styled(doc, "II.2. Kiến trúc RAG hiện đại", level=2)
    doc.add_paragraph(
        "Quy trình RAG hiện đại gồm các bước:\n\n"
        "A. Offline Indexing:\n"
        "   1. Document Loading: parse PDF/DOCX/TXT, trích xuất text + metadata\n"
        "   2. Text Cleaning: chuẩn hóa Unicode NFC, loại bỏ ký tự nhiễu\n"
        "   3. Chunking: chia nhỏ văn bản (512 tokens, overlap 128)\n"
        "   4. Embedding: encode mỗi chunk thành vector (dense 1024-d + sparse)\n"
        "   5. Indexing: upsert vectors vào Qdrant\n\n"
        "B. Online Query:\n"
        "   1. Query Encoding: embed câu hỏi thành vector\n"
        "   2. Hybrid Search: dense ANN + sparse BM25-like\n"
        "   3. RRF Fusion: gộp kết quả từ 2 nguồn\n"
        "   4. Cross-Encoder Reranking: đánh giá lại top-30 → top-5\n"
        "   5. LLM Generation: sinh câu trả lời có citation"
    )

    add_heading_styled(doc, "II.3. Hybrid Search + Reranking", level=2)
    doc.add_paragraph(
        "Hybrid Search kết hợp ưu điểm của 2 phương pháp:\n"
        "• Dense (semantic): bắt được ngữ nghĩa, câu hỏi tự nhiên\n"
        "• Sparse (keyword): chính xác với tên riêng, số hiệu, từ khóa kỹ thuật\n\n"
        "RRF Fusion (k=60): score(d) = Σ 1/(k + rank_i + 1) cho mỗi ranked list i\n\n"
        "Cross-Encoder Reranking:\n"
        "• Bi-encoder (retrieve): encode query và doc độc lập → nhanh nhưng thô\n"
        "• Cross-encoder (rerank): đọc cả query + doc cùng lúc → chậm hơn nhưng "
        "chính xác hơn nhiều\n"
        "• Pipeline hình phễu: Retrieve 30 (fast) → Rerank top-5 (accurate)"
    )
    doc.add_page_break()

    # === III. KIẾN TRÚC HỆ THỐNG ===
    add_heading_styled(doc, "III. Kiến trúc hệ thống", level=1)

    add_heading_styled(doc, "III.1. Tổng quan kiến trúc", level=2)
    doc.add_paragraph(
        "Hệ thống gồm 3 tầng chính:\n\n"
        "┌─────────────────────────────────────────────┐\n"
        "│  Frontend: Next.js 14 (port 3000)           │\n"
        "│  Chat UI + PDF viewer + Image search        │\n"
        "└──────────────────┬──────────────────────────┘\n"
        "                   │ HTTP / SSE\n"
        "┌──────────────────▼──────────────────────────┐\n"
        "│  Backend: FastAPI (port 8000)                │\n"
        "│  /chat/stream, /image-search, /file         │\n"
        "└───────┬──────────────────┬──────────────────┘\n"
        "        │                  │\n"
        "┌───────▼───────┐  ┌──────▼──────────────────┐\n"
        "│ Qdrant (local) │  │ Ollama (Qwen3-4B, GPU) │\n"
        "│ hr_documents   │  │ bge-m3, bge-reranker   │\n"
        "│ hr_images      │  │ CLIP ViT-B/32          │\n"
        "└────────────────┘  └─────────────────────────┘"
    )

    add_heading_styled(doc, "III.2. Ingestion Pipeline", level=2)
    doc.add_paragraph(
        "File (PDF/DOCX/TXT)\n"
        "  → parse_file() [PyMuPDF / python-docx]\n"
        "  → clean_text() [NFC normalize, strip control chars]\n"
        "  → chunk_documents() [RecursiveCharacterTextSplitter, 512/128]\n"
        "  → embedder.encode_batch() [bge-m3: dense 1024-d + sparse]\n"
        "  → Qdrant upsert (collection: hr_documents)\n\n"
        "Song song:\n"
        "  → extract_images() [PyMuPDF: trích ảnh embedded]\n"
        "  → clip.encode_image() [CLIP ViT-B/32: 512-d]\n"
        "  → Qdrant upsert (collection: hr_images)\n\n"
        "Incremental: manifest SHA256 theo dõi file đã index, chỉ re-embed khi nội dung thay đổi."
    )

    add_heading_styled(doc, "III.3. Search + Generation Pipeline", level=2)
    doc.add_paragraph(
        "Query: 'Cấu trúc bài thi VSTEP phần Speaking?'\n"
        "  → bge-m3 encode → q_dense (1024-d) + q_sparse\n"
        "  → Qdrant query_batch_points (dense + sparse, limit=30)\n"
        "  → RRF fusion → top 30 fused\n"
        "  → bge-reranker-v2-m3 cross-encoder → top 5\n"
        "  → Filter: loại chunks có score < 0.35\n"
        "  → Dedup: gộp chunks cùng file + cùng trang\n"
        "  → Build RAG prompt (system + context + question)\n"
        "  → Ollama Qwen3-4B stream (num_gpu=99, num_ctx=4096)\n"
        "  → Filter <think>...</think> tags\n"
        "  → SSE stream tokens → Frontend render real-time"
    )

    add_heading_styled(doc, "III.4. Image Search (CLIP)", level=2)
    doc.add_paragraph(
        "User gửi ảnh vào chat:\n"
        "  → CLIP ViT-B/32 encode image → q_img (512-d)\n"
        "  → Qdrant search collection hr_images (cosine similarity)\n"
        "  → Top-5 ảnh tương tự + metadata (file gốc, trang, caption)\n"
        "  → Frontend render DocCard với PDF viewer mở đúng trang\n\n"
        "CLIP cho phép tìm ảnh bằng ảnh (reverse image search) hoặc bằng text "
        "(cross-modal) vì cả 2 modality nằm trong cùng không gian 512-d."
    )
    doc.add_page_break()

    # === IV. CÔNG NGHỆ ===
    add_heading_styled(doc, "IV. Công nghệ sử dụng", level=1)

    add_heading_styled(doc, "IV.1. Bảng công nghệ", level=2)
    tech_rows = [
        ("LLM", "Ollama + Qwen3-4B (Q4_K_M)", "Sinh câu trả lời từ context"),
        ("Embedding", "BAAI/bge-m3", "Dense (1024-d) + Sparse vectors"),
        ("Reranker", "BAAI/bge-reranker-v2-m3", "Cross-encoder rerank top-30 → top-5"),
        ("Image Search", "CLIP ViT-B/32 multilingual", "Reverse image search 512-d"),
        ("Vector DB", "Qdrant (embedded, local)", "Hybrid search, no Docker needed"),
        ("Backend", "FastAPI + SSE streaming", "REST API + real-time token stream"),
        ("Frontend", "Next.js 14 + Tailwind + shadcn", "Chat UI + inline PDF viewer"),
        ("PDF Parser", "PyMuPDF (fitz)", "Extract text + images per page"),
        ("DOCX Parser", "python-docx + mammoth", "Parse + HTML preview"),
        ("Text Splitter", "langchain-text-splitters", "Recursive chunking 512/128"),
    ]
    add_table(doc, ["Thành phần", "Công nghệ", "Vai trò"], tech_rows)

    add_heading_styled(doc, "IV.2. Giải thích lựa chọn model", level=2)
    doc.add_paragraph(
        "1. bge-m3 (Embedding):\n"
        "   • Output cả dense + sparse trong 1 model → giảm memory, đồng nhất semantic space\n"
        "   • Support 100+ ngôn ngữ, context 8K tokens\n"
        "   • Top MTEB multilingual benchmark\n\n"
        "2. bge-reranker-v2-m3 (Reranker):\n"
        "   • Cross-encoder đọc query + doc đồng thời → bắt quan hệ tinh tế\n"
        "   • Multilingual, 568M params, chạy được trên 8GB VRAM\n\n"
        "3. Qwen3-4B (LLM):\n"
        "   • 4B params Q4_K_M ≈ 2.6GB VRAM → fit cùng embedder + reranker trên 8GB\n"
        "   • Multilingual (Vietnamese tốt), Apache 2.0 license\n"
        "   • Thinking mode toggle (đã disable để tránh empty response)\n\n"
        "4. CLIP ViT-B/32:\n"
        "   • Multilingual text encoder + image encoder cùng 512-d space\n"
        "   • ~600MB, lazy load on-demand\n"
        "   • Cho phép tìm ảnh bằng text hoặc bằng ảnh\n\n"
        "5. Qdrant (Vector DB):\n"
        "   • Viết bằng Rust → tốc độ cao\n"
        "   • Native sparse vectors cho hybrid search\n"
        "   • Embedded mode (no Docker) cho local deployment"
    )
    doc.add_page_break()

    # === V. GIAO DIỆN ===
    add_heading_styled(doc, "V. Giao diện người dùng", level=1)
    doc.add_paragraph(
        "Frontend xây dựng bằng Next.js 14 App Router + Tailwind CSS + shadcn-style components.\n\n"
        "Các tính năng chính:\n"
        "• Chat interface với streaming token (SSE)\n"
        "• Inline PDF viewer (iframe, mở đúng trang kết quả)\n"
        "• DOCX/TXT preview (convert sang HTML on-the-fly)\n"
        "• Image search: drag-drop / paste ảnh vào chat\n"
        "• Nút 📂 mở file gốc trên máy (os.startfile)\n"
        "• Nút ✕ đóng từng card kết quả\n"
        "• Typing indicator (bouncing dots) khi đang xử lý\n"
        "• Welcome screen với 4 example prompts\n"
        "• Responsive, theme trắng chuyên nghiệp"
    )
    doc.add_page_break()

    # === VI. KẾT QUẢ ===
    add_heading_styled(doc, "VI. Kết quả và đánh giá", level=1)

    doc.add_paragraph("Hiệu năng đo trên RTX 5060 (8GB VRAM), Qwen3-4B, 39 tài liệu:")
    perf_rows = [
        ("Search (hybrid)", "200-800 ms", "30 chunks retrieved"),
        ("Rerank (cross-encoder)", "150-300 ms", "Top-5 from 30"),
        ("Generate (streaming)", "3-15 s", "Qwen3-4B, 100% GPU"),
        ("Image search (CLIP)", "30-1000 ms", "5 hits, first load ~6s"),
        ("Total /chat latency", "5-20 s", "End-to-end"),
        ("VRAM peak", "~7.5 GB", "All models loaded"),
    ]
    add_table(doc, ["Stage", "Thời gian", "Ghi chú"], perf_rows)

    doc.add_paragraph("\nĐộ chính xác (đánh giá thủ công trên 20 câu hỏi):")
    acc_rows = [
        ("Retrieval Recall@5", "~85%", "Tài liệu đúng nằm trong top-5"),
        ("Answer Accuracy", "~75%", "Câu trả lời chứa thông tin đúng"),
        ("Image Search Precision", "~80%", "Ảnh trả về liên quan đến query"),
    ]
    add_table(doc, ["Metric", "Giá trị", "Mô tả"], acc_rows)
    doc.add_page_break()

    # === VII. KẾT LUẬN ===
    add_heading_styled(doc, "VII. Kết luận và hướng phát triển", level=1)
    doc.add_paragraph(
        "Hệ thống Document Search Assistant đã đạt được các mục tiêu đề ra:\n"
        "• Chạy 100% local, dữ liệu không rời máy\n"
        "• Hybrid search + reranking cho kết quả chính xác\n"
        "• Image search bằng CLIP mở rộng khả năng tìm kiếm\n"
        "• Streaming response giảm perceived latency\n"
        "• Giao diện Next.js hiện đại, dễ sử dụng\n\n"
        "Hướng phát triển:\n"
        "• Multi-query: LLM sinh 3 biến thể câu hỏi → tăng recall\n"
        "• Parent-child chunking: search child (precise), return parent (rich context)\n"
        "• Conversation memory: hỗ trợ follow-up questions\n"
        "• OCR cho scanned PDF (VietOCR)\n"
        "• Streaming response nhanh hơn với vLLM thay Ollama\n"
        "• Deploy multi-user với Qdrant server mode"
    )

    # Save
    doc.save(str(OUTPUT_FILE))
    print(f"Report saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_report()
