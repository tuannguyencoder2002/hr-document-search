"""Parse PDF / DOCX / TXT into per-page documents with metadata."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class ParsedDocument:
    """A single parsed unit (typically one page)."""

    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def _ensure_file(path: str | Path) -> Path:
    p = Path(path)
    if not p.exists():
        raise ValueError(f"File not found: {p}")
    if not p.is_file():
        raise ValueError(f"Not a file: {p}")
    return p


def parse_pdf(path: str | Path) -> list[ParsedDocument]:
    """Parse a PDF file into a list of per-page ParsedDocument."""
    import fitz  # PyMuPDF

    p = _ensure_file(path)
    if p.suffix.lower() != ".pdf":
        raise ValueError(f"Expected .pdf, got {p.suffix}")

    docs: list[ParsedDocument] = []
    with fitz.open(p) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            text = page.get_text("text") or ""
            if not text.strip():
                continue
            docs.append(
                ParsedDocument(
                    text=text,
                    metadata={
                        "source": p.name,
                        "source_path": str(p),
                        "page": page_index,
                        "file_type": "pdf",
                    },
                )
            )
    return docs


def parse_docx(path: str | Path) -> list[ParsedDocument]:
    """Parse a DOCX file. DOCX has no real pages; treat the whole file as one doc."""
    from docx import Document

    p = _ensure_file(path)
    if p.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx, got {p.suffix}")

    doc = Document(str(p))
    parts: list[str] = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        return []
    return [
        ParsedDocument(
            text=text,
            metadata={
                "source": p.name,
                "source_path": str(p),
                "page": 1,
                "file_type": "docx",
            },
        )
    ]


def parse_txt(path: str | Path) -> list[ParsedDocument]:
    """Parse a plain text file."""
    p = _ensure_file(path)
    text = p.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    return [
        ParsedDocument(
            text=text,
            metadata={
                "source": p.name,
                "source_path": str(p),
                "page": 1,
                "file_type": "txt",
            },
        )
    ]


def parse_file(path: str | Path) -> list[ParsedDocument]:
    """Dispatch parser based on file extension."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(p)
    if ext == ".docx":
        return parse_docx(p)
    if ext in {".txt", ".md"}:
        return parse_txt(p)
    raise ValueError(f"Unsupported file extension: {ext}")
